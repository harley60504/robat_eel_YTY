from __future__ import annotations

import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "outputs" / "algorithm_doc_assets"
DOCX_PATH = ROOT / "outputs" / "eel_cpg_path_following_algorithm.docx"
GAIT_DIR = ROOT / "gaits"
GAIT_ORDER = [
    "straight.json",
    "turn_left.json",
    "turn_right.json",
    "spin_left.json",
    "spin_right.json",
]


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/msjhbd.ttc" if bold else "C:/Windows/Fonts/msjh.ttc",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_round_box(draw, xy, text, fill, outline="#31506B", text_color="#111827"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    lines = text.split("\n")
    f = font(24, bold=True)
    total_h = sum(draw.textbbox((0, 0), line, font=f)[3] for line in lines) + (len(lines) - 1) * 8
    y = y1 + ((y2 - y1) - total_h) / 2 - 3
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=f)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), line, font=f, fill=text_color)
        y += (bbox[3] - bbox[1]) + 8


def arrow(draw, start, end, color="#334155"):
    draw.line([start, end], fill=color, width=4)
    import math

    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 14
    p1 = (end[0] - size * math.cos(angle - 0.45), end[1] - size * math.sin(angle - 0.45))
    p2 = (end[0] - size * math.cos(angle + 0.45), end[1] - size * math.sin(angle + 0.45))
    draw.polygon([end, p1, p2], fill=color)


def make_overall_flow(path: Path):
    img = Image.new("RGB", (1500, 760), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(34, bold=True)
    draw.text((50, 32), "階層式 CPG 路徑追蹤控制架構", font=title_font, fill="#0F172A")

    boxes = [
        ((70, 135, 365, 255), "Camera / MuJoCo\n取得 x, y, yaw", "#E0F2FE"),
        ((455, 135, 750, 255), "Pure Pursuit\n前視點路徑追蹤", "#DCFCE7"),
        ((840, 135, 1135, 255), "Steering Law\nheading error → steer", "#FEF3C7"),
        ((455, 395, 750, 515), "Hopf CPG\n產生 traveling wave", "#EDE9FE"),
        ((840, 395, 1135, 515), "Joint Bias Mapping\nsteer → joint_bias_i", "#FCE7F3"),
        ((1215, 265, 1455, 385), "Servo Command\nq_i(t)", "#F1F5F9"),
    ]
    for b in boxes:
        draw_round_box(draw, *b)

    arrow(draw, (365, 195), (455, 195))
    arrow(draw, (750, 195), (840, 195))
    arrow(draw, (1135, 195), (1270, 265))
    arrow(draw, (750, 455), (840, 455))
    arrow(draw, (1135, 455), (1270, 385))
    arrow(draw, (575, 255), (575, 395))
    arrow(draw, (980, 255), (980, 395))

    small = font(22)
    draw.text((82, 600), "輸出公式：q_i(t) = A_joint · r_i · cos(theta_i) + w_i · u_turn", font=font(28, bold=True), fill="#1E293B")
    draw.text((82, 650), "底層 CPG 提供穩定節律；上層 pure pursuit 依據路徑誤差產生轉向偏置。", font=small, fill="#334155")
    img.save(path)


def make_physical_arch(path: Path):
    img = Image.new("RGB", (1500, 720), "white")
    draw = ImageDraw.Draw(img)
    draw.text((50, 30), "實體控制架構比較", font=font(34, bold=True), fill="#0F172A")

    draw.text((85, 110), "Mode A：PC Angle Mode", font=font(28, bold=True), fill="#1D4ED8")
    y = 170
    boxes_a = [
        ((80, y, 360, y + 95), "PC / Python\nCPG + Controller", "#DBEAFE"),
        ((430, y, 710, y + 95), "Serial\n6 servo angles", "#E0F2FE"),
        ((780, y, 1060, y + 95), "Control Board\nwrite servo", "#F1F5F9"),
        ((1130, y, 1410, y + 95), "Servo Motors\njoint angles", "#F8FAFC"),
    ]
    for b in boxes_a:
        draw_round_box(draw, *b)
    for s, e in [((360, y + 48), (430, y + 48)), ((710, y + 48), (780, y + 48)), ((1060, y + 48), (1130, y + 48))]:
        arrow(draw, s, e)

    draw.text((85, 385), "Mode B：On-board CPG Mode", font=font(28, bold=True), fill="#047857")
    y = 445
    boxes_b = [
        ((80, y, 360, y + 95), "PC / Python\nCamera + Pure Pursuit", "#DCFCE7"),
        ((430, y, 710, y + 95), "Serial\nsteer / gait params", "#ECFDF5"),
        ((780, y, 1060, y + 95), "Control Board\nHopf CPG", "#F1F5F9"),
        ((1130, y, 1410, y + 95), "Servo Motors\njoint angles", "#F8FAFC"),
    ]
    for b in boxes_b:
        draw_round_box(draw, *b)
    for s, e in [((360, y + 48), (430, y + 48)), ((710, y + 48), (780, y + 48)), ((1060, y + 48), (1130, y + 48))]:
        arrow(draw, s, e)

    img.save(path)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def fmt(values, digits: int = 3) -> str:
    return ", ".join(f"{float(value):.{digits}f}" for value in values)


def load_gaits() -> list[dict]:
    gaits = []
    for name in GAIT_ORDER:
        path = GAIT_DIR / name
        with path.open("r", encoding="utf-8") as f:
            gait = json.load(f)
        gait["file"] = name
        gaits.append(gait)
    return gaits


def amp_scales_to_mu_scales(amp_scales: tuple[float, ...]) -> tuple[float, ...]:
    return tuple(float(value) * float(value) for value in amp_scales)


def turning_amp_scales(base_scales: tuple[float, ...], steer: float, gain: float) -> tuple[float, ...]:
    if gain <= 0.0:
        return base_scales
    tail_weights = (0.0, 0.0, 0.15, 0.35, 0.65, 1.0)
    values = []
    for base, weight in zip(base_scales, tail_weights):
        multiplier = 1.0 + gain * abs(float(steer)) * weight
        values.append(min(1.6, max(0.2, float(base) * multiplier)))
    return tuple(values)


def add_table(doc: Document, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    return table


def set_doc_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(31, 78, 121)),
    ]:
        style = styles[name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def add_formula_block(doc: Document, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    cell.text = ""
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(10)
    return table


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    overall_png = OUT_DIR / "overall_flow.png"
    arch_png = OUT_DIR / "physical_architecture.png"
    make_overall_flow(overall_png)
    make_physical_arch(arch_png)
    gaits = load_gaits()

    doc = Document()
    set_doc_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("機器鰻 Hopf CPG 與矩形路徑追蹤控制演算法")
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("CPG traveling wave、轉向偏置、Pure Pursuit 與實體控制架構整理")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_heading("1. 控制架構總覽", level=1)
    doc.add_paragraph(
        "本控制系統採用階層式架構：底層 Hopf CPG 產生穩定的身體 traveling wave，"
        "上層路徑追蹤器根據機器魚位置與朝向計算轉向命令，並將轉向命令映射為各關節偏置角。"
        "因此，游動節律與路徑導航可以分開設計，便於調參、實體移植與論文分析。"
    )
    doc.add_picture(str(overall_png), width=Inches(6.6))

    doc.add_heading("2. Hopf CPG 公式", level=1)
    doc.add_paragraph("每一節關節對應一個 Hopf oscillator，其狀態包含振幅 r_i 與相位 theta_i。")
    add_formula_block(
        doc,
        [
            "r_dot_i = alpha (mu - r_i^2) r_i + k_fb_amp fb_amp",
            "theta_dot_i = omega + coupling_i + anchor_i + k_fb_phase fb_phase",
            "omega = 2 pi f",
            "mu_i = amp_i^2",
            "q_i(t) = A_joint r_i cos(theta_i) + b_i",
        ],
    )
    doc.add_paragraph(
        "其中 r_dot_i 控制振幅收斂，theta_dot_i 控制相位前進。"
        "當 fb_amp 與 fb_phase 設為 0 時，CPG 主要由基本頻率、相鄰耦合與參考相位牽引所決定。"
    )

    add_table(
        doc,
        ["符號", "意義", "目前用途"],
        [
            ["r_i", "第 i 節 oscillator 的振幅狀態", "由 Hopf 方程收斂至穩態振幅"],
            ["theta_i", "第 i 節 oscillator 的相位狀態", "決定 servo 在週期中的位置"],
            ["alpha", "振幅收斂速度", "越大代表啟動時振幅越快收斂"],
            ["mu", "穩態振幅平方", "mu=1 時 r_i 收斂至 1"],
            ["omega", "基本角速度", "omega=2 pi f"],
            ["q_i(t)", "第 i 節 servo 目標角", "CPG 波形加上轉向偏置"],
        ],
        [1.05, 2.55, 2.7],
    )

    doc.add_heading("3. 相位耦合與 Anchor 項", level=1)
    doc.add_paragraph(
        "相位耦合項 coupling_i 維持相鄰 oscillator 的目標相位差，使各關節形成穩定 traveling wave。"
        "anchor_i 則將每一節 oscillator 弱牽引至全域參考波，避免整體相位漂移，並提升 reset 或擾動後的收斂一致性。"
    )
    add_formula_block(
        doc,
        [
            "coupling_i = K_couple sin((theta_{i-1} - theta_i) - desired_delta_{i-1,i})",
            "           + K_couple sin((theta_{i+1} - theta_i) - desired_delta_{i+1,i})",
            "theta_ref_i = omega t + phase_offset_i",
            "anchor_i = K_anchor sin(theta_ref_i - theta_i)",
        ],
    )

    doc.add_heading("4. 外層調參：amp_scales、phase_lags、joint_bias", level=1)
    add_table(
        doc,
        ["參數", "演算法角色", "物理意義"],
        [
            ["amp_scales", "使用者輸入的各節目標振幅倍率，內部轉成 mu_i = amp_i^2", "調整各節擺幅包絡，例如頭部或尾部加大"],
            ["phase_lags", "決定 phase_offset_i", "控制相鄰節相位差與等效波長"],
            ["joint_bias b_i", "直接加在 q_i 上", "讓身體平均彎向一側以產生轉彎"],
            ["steering bias u_turn", "由 pure pursuit 產生", "根據路徑誤差即時控制轉向大小與方向"],
        ],
        [1.5, 2.35, 2.45],
    )
    add_formula_block(
        doc,
        [
            "b_i = w_i u_turn",
            "w = [0.45, 0.55, 0.68, 0.80, 0.92, 1.00]",
            "u_turn,t = u_turn,t-1 + alpha_s (u_target - u_turn,t-1)",
        ],
    )

    doc.add_heading("5. 轉彎時的振幅調變", level=1)
    doc.add_paragraph(
        "為提升轉彎能力，本研究在轉彎時加入尾部振幅調變。其概念是在 steering command 絕對值變大時，"
        "提高後段關節的 Hopf 振幅目標，使尾部在轉彎時產生更強的橫向推進與偏航力矩。此方法不改變 CPG 的相位結構，"
        "而是將目標振幅轉換為各關節 mu_i，使振幅狀態 r_i 透過 Hopf 方程平滑收斂。"
    )
    add_formula_block(
        doc,
        [
            "amp_i,turn = amp_i [1 + k_amp |u_turn| gamma_i]",
            "mu_i,turn = (amp_i,turn)^2",
            "gamma = [0.00, 0.00, 0.15, 0.35, 0.65, 1.00]",
            "r_dot_i = alpha (mu_i,turn - r_i^2) r_i",
            "q_i(t) = A_joint r_i cos(theta_i) + w_i u_turn",
            "current default: k_amp = 0.6",
        ],
    )
    add_table(
        doc,
        ["參數", "意義", "目前設定"],
        [
            ["k_amp / turn_amp_gain", "轉彎振幅目標增益；越大代表轉彎時尾部目標振幅增加越多", "0.6"],
            ["gamma_i", "各節振幅調變權重，前段小、尾段大", "[0, 0, 0.15, 0.35, 0.65, 1.0]"],
            ["|u_turn|", "轉向命令絕對值", "由 pure pursuit heading error 計算"],
            ["mu_i,turn", "轉彎時 Hopf oscillator 的振幅目標平方", "由 amp_i,turn 平方取得"],
        ],
        [1.5, 3.05, 1.75],
    )
    doc.add_paragraph(
        "固定轉彎測試顯示，尾部振幅調變可提高 yaw rate 並降低轉彎半徑。"
        "因此矩形環繞控制目前預設啟用 k_amp = 0.6，以取得較佳轉彎能力與穩定性折衷。"
    )
    add_table(
        doc,
        ["測試條件", "turn_amp_gain", "mean yaw rate", "estimated turn radius"],
        [
            ["中等轉彎 bias = [0.08, ..., 0.18]", "0.0", "-0.312 rad/s", "0.936 m"],
            ["中等轉彎 bias = [0.08, ..., 0.18]", "0.6", "-0.345 rad/s", "0.863 m"],
            ["中等轉彎 bias = [0.08, ..., 0.18]", "1.0", "-0.372 rad/s", "0.810 m"],
            ["小圈 bias = [0.12, ..., 0.27]", "0.0", "-0.459 rad/s", "0.658 m"],
            ["小圈 bias = [0.12, ..., 0.27]", "0.6", "-0.539 rad/s", "0.560 m"],
            ["小圈 bias = [0.12, ..., 0.27]", "1.0", "-0.556 rad/s", "0.544 m"],
        ],
        [2.55, 1.2, 1.3, 1.35],
    )

    doc.add_heading("6. 直線與轉彎 Gait 參數表", level=1)
    doc.add_paragraph(
        "目前固定 gait 檔案位於 gaits 資料夾，包含直線、左右轉彎與左右原地轉圈。"
        "這些 gait 使用相同的基礎頻率、波長、相位差與振幅包絡，主要差異在 joint_bias。"
        "joint_bias 為正時身體平均彎向一側；joint_bias 為負時產生鏡像方向的轉彎。"
    )
    add_table(
        doc,
        ["Gait", "檔案", "A_joint", "f", "lambda", "amp_scales", "phase_lags", "joint_bias"],
        [
            [
                gait["name"],
                gait["file"],
                f'{float(gait["ajoint"]):.3f}',
                f'{float(gait["freq"]):.3f}',
                f'{float(gait["wavelength"]):.4f}',
                fmt(gait["amp_scales"]),
                fmt(gait["phase_lags"], 6),
                fmt(gait["joint_bias"]),
            ]
            for gait in gaits
        ],
        [0.9, 1.0, 0.75, 0.65, 0.8, 1.75, 1.55, 1.55],
    )
    straight = next(gait for gait in gaits if gait["name"] == "straight")
    doc.add_paragraph(
        "直線游動的 joint_bias 全部為 0，因此輸出為對稱 traveling wave。"
        "目前直線 gait 的 amp_scales 已經包含尾部加大："
        f"{fmt(straight['amp_scales'])}。"
        "也就是說，直線游時尾部振幅不是 1.0，而是已經提高到 1.20；"
        "轉彎控制則是在此基礎上，依 steering command 對後段振幅再額外放大。"
    )
    add_table(
        doc,
        ["Gait 類型", "控制方式", "主要作用"],
        [
            ["straight", "joint_bias = [0, 0, 0, 0, 0, 0]", "對稱身體波，主要追求前進速度"],
            ["turn_left / turn_right", "joint_bias 線性增加並正負鏡像", "形成穩定轉彎半徑"],
            ["spin_left / spin_right", "較大的 joint_bias 正負鏡像", "產生較大的 yaw rate，接近原地繞圈"],
        ],
        [1.25, 2.4, 2.6],
    )

    doc.add_heading("7. 矩形環繞中的動態 Gait 生成", level=1)
    doc.add_paragraph(
        "矩形環繞並不是硬切換 straight、turn_left 或 spin_left 檔案，而是由 pure pursuit 連續產生 steer。"
        "steer 經低通濾波後轉成 joint_bias，同時依 |steer| 調整尾段 mu_i，因此轉彎可以在直線 gait 與轉彎 gait 之間平滑過渡。"
    )
    add_formula_block(
        doc,
        [
            "steer_state_t = steer_state_{t-1} + beta (steer_target - steer_state_{t-1})",
            "joint_bias_i = w_i steer_state_t",
            "amp_i,turn = amp_i [1 + k_amp |steer_state_t| gamma_i]",
            "mu_i,turn = (amp_i,turn)^2",
        ],
    )
    base_scales = tuple(float(value) for value in straight["amp_scales"])
    example_steer = 0.18
    turn_gain = 0.6
    turned_scales = turning_amp_scales(base_scales, example_steer, turn_gain)
    add_table(
        doc,
        ["案例", "amp_scales", "mu_scales"],
        [
            ["直線基礎值", fmt(base_scales), fmt(amp_scales_to_mu_scales(base_scales))],
            [
                f"轉彎範例 steer={example_steer:.2f}, k_amp={turn_gain:.1f}",
                fmt(turned_scales),
                fmt(amp_scales_to_mu_scales(turned_scales)),
            ],
        ],
        [1.8, 2.45, 2.2],
    )
    doc.add_paragraph(
        "目前 Hopf CPG 的 alpha=4.0、K_couple=0.35、K_anchor=0.10，"
        "相較於較快的設定，切換振幅或轉向命令時會更慢收斂，但能減少游法切換時的瞬間卡頓。"
    )

    doc.add_heading("8. Pure Pursuit 路徑追蹤", level=1)
    doc.add_paragraph(
        "矩形環繞控制採用 pure pursuit，而非只追四個角點。演算法先將機器魚目前位置投影到矩形路徑，"
        "再沿路徑往前取 lookahead distance 的前視點。魚體轉向目標為該前視點，因此轉彎會提前且連續，"
        "避免 waypoint 硬切換造成卡頓。"
    )
    add_formula_block(
        doc,
        [
            "target = rectangle_path.point_at(closest_s(x, y) + lookahead)",
            "desired_yaw = atan2(target_y - y, target_x - x)",
            "heading_error = wrap_pi(desired_yaw - yaw)",
            "u_target = clip(-K heading_error, -u_max, u_max)",
        ],
    )
    add_table(
        doc,
        ["參數", "目前值", "效果"],
        [
            ["path-half-x", "1.10 m", "矩形追蹤路徑的半長"],
            ["path-half-y", "0.35 m", "矩形追蹤路徑的半寬"],
            ["lookahead", "0.50 m", "越大越早轉且更平滑，越小越貼近路徑但容易急轉"],
            ["steer_gain K", "0.55", "heading error 轉為 steering 的比例"],
            ["max_bias", "0.34 rad", "限制最大轉向偏置，避免關節過度彎曲"],
            ["steer_smoothing", "0.08", "轉向低通濾波，降低 waypoint/路徑切換造成的突變"],
        ],
        [1.35, 1.1, 3.85],
    )

    doc.add_heading("9. 實體控制架構：兩種模式", level=1)
    doc.add_picture(str(arch_png), width=Inches(6.6))
    add_table(
        doc,
        ["模式", "電腦端", "控制板端", "優點", "限制"],
        [
            [
                "PC Angle Mode",
                "計算 CPG、路徑控制與 6 顆 servo 角度",
                "接收角度並輸出 servo",
                "演算法容易修改，MuJoCo 與實體較易一致",
                "通訊延遲與 Python 卡頓會直接影響 servo 平順性",
            ],
            [
                "On-board CPG Mode",
                "以 camera 估測 x,y,yaw，計算 steer 或 gait 參數",
                "在控制板上更新 CPG 並輸出 servo",
                "servo 更新頻率穩定，通訊需求低，較適合長時間運行",
                "需確保控制板 CPG 與 Python/MuJoCo 參數一致",
            ],
        ],
        [1.1, 1.45, 1.35, 1.35, 1.25],
    )

    doc.add_heading("10. 可放入論文的方法描述", level=1)
    doc.add_paragraph(
        "本研究採用階層式 CPG 路徑追蹤控制架構。底層以 Hopf CPG 產生週期性身體波動，"
        "上層使用 pure pursuit 方法根據機器魚當前位置與目標路徑計算轉向命令。"
        "該轉向命令依各關節權重轉換為關節偏置角，並疊加於 CPG 輸出，以達成矩形路徑追蹤。"
    )
    doc.add_paragraph(
        "此方法保留 CPG 的節律穩定性，同時透過 amplitude scaling、phase-lag specification 與 steering bias "
        "提供任務層級的可控性。相較於直接以正弦波命令所有關節，Hopf CPG 具有振幅收斂、相位同步與回饋整合能力；"
        "相較於端到端強化學習，該方法的參數具備明確物理意義，便於實體調校與結果分析。"
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Robot eel MuJoCo control algorithm summary"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from hopf_cpg import HopfCPGParams


ROOT = Path(__file__).resolve().parent
GAIT_DIR = ROOT / "gaits"
OUT = ROOT / "outputs" / "eel_straight_turning_gaits.docx"


GAIT_ORDER = [
    "straight.json",
    "turn_left.json",
    "turn_right.json",
    "spin_left.json",
    "spin_right.json",
]


def load_gaits() -> list[dict]:
    gaits = []
    for name in GAIT_ORDER:
        path = GAIT_DIR / name
        with path.open("r", encoding="utf-8") as f:
            gait = json.load(f)
        gait["file"] = name
        gaits.append(gait)
    return gaits


def fmt(values, digits: int = 3) -> str:
    return ", ".join(f"{float(value):.{digits}f}" for value in values)


def turning_amp_scales(base_scales: tuple[float, ...], steer: float, gain: float) -> tuple[float, ...]:
    if gain <= 0.0:
        return base_scales
    tail_weights = (0.0, 0.0, 0.15, 0.35, 0.65, 1.0)
    values = []
    for base, weight in zip(base_scales, tail_weights):
        multiplier = 1.0 + gain * abs(float(steer)) * weight
        values.append(min(1.6, max(0.2, base * multiplier)))
    return tuple(float(value) for value in values)


def amp_scales_to_mu_scales(amp_scales: tuple[float, ...]) -> tuple[float, ...]:
    return tuple(float(value) * float(value) for value in amp_scales)


def add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers):
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    doc.add_paragraph()
    return table


def build():
    gaits = load_gaits()
    params = HopfCPGParams()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_heading("Robot Eel Straight and Turning Gait Parameters", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Purpose. ").bold = True
    p.add_run(
        "This document summarizes the currently used fixed straight, turning, and spin gaits "
        "for the MuJoCo robot eel model. The gait generator is a Hopf CPG with per-joint "
        "amplitude targets represented internally by mu_i, phase lag, and steering bias."
    )

    add_heading(doc, "1. Hopf CPG Used in the Simulation", 1)
    doc.add_paragraph("For joint i, the controller uses the following oscillator form:")
    doc.add_paragraph("r_dot_i = alpha (mu_i - r_i^2) r_i + k_fb_amp fb_amp")
    doc.add_paragraph("theta_dot_i = omega + coupling_i + anchor_i + k_fb_phase fb_phase")
    doc.add_paragraph("mu_i = amp_i^2")
    doc.add_paragraph("q_i = A_joint r_i cos(theta_i) + b_i")
    doc.add_paragraph(
        "where q_i is the target joint angle, amp_i is the human-readable per-joint "
        "amplitude scale, mu_i is the Hopf amplitude target, b_i is the joint steering "
        "bias, and omega = 2 pi f."
    )

    add_table(
        doc,
        ["Parameter", "Current value", "Meaning"],
        [
            ["A_joint", "0.45 rad", "Base joint amplitude"],
            ["frequency", "1.0 Hz", "Oscillation frequency"],
            ["wavelength", "1.6275", "Nominal body wave parameter"],
            ["phase lag", "0.614439 rad/joint", "Phase difference between adjacent joints"],
            ["alpha", f"{params.alpha:.2f}", "Hopf amplitude convergence speed"],
            ["k_couple", f"{params.k_couple:.2f}", "Adjacent-joint phase synchronization gain"],
            ["k_anchor", f"{params.k_anchor:.2f}", "Reference wave anchoring gain"],
        ],
    )

    add_heading(doc, "2. Fixed Gait Parameter Table", 1)
    rows = []
    for gait in gaits:
        rows.append(
            [
                gait["name"],
                gait["file"],
                f'{float(gait["ajoint"]):.3f}',
                f'{float(gait["freq"]):.3f}',
                f'{float(gait["wavelength"]):.4f}',
                fmt(gait["amp_scales"]),
                fmt(gait["phase_lags"], 6),
                fmt(gait["joint_bias"]),
            ]
        )
    add_table(
        doc,
        [
            "Gait",
            "File",
            "A_joint",
            "f (Hz)",
            "lambda",
            "amp_scales",
            "phase_lags",
            "joint_bias",
        ],
        rows,
    )

    add_heading(doc, "3. Straight Swimming Gait", 1)
    straight = next(g for g in gaits if g["name"] == "straight")
    doc.add_paragraph(
        "The straight gait uses zero joint bias, so the body wave is symmetric and the eel "
        "mainly generates forward thrust. The current straight-swim amplitude pattern is:"
    )
    doc.add_paragraph(fmt(straight["amp_scales"]))
    doc.add_paragraph(
        "This pattern already increases the head-side and tail-side amplitudes relative to the "
        "middle joints. Therefore, the tail amplitude is already larger during straight swimming."
    )

    add_heading(doc, "4. Turning and Spin Gaits", 1)
    doc.add_paragraph(
        "Turning is produced by adding a static joint_bias to the CPG output. Positive bias bends "
        "the body to one side, while negative bias mirrors the gait in the opposite direction. "
        "The phase lags are kept fixed for these saved gaits."
    )
    add_table(
        doc,
        ["Gait", "joint_bias"],
        [[gait["name"], fmt(gait["joint_bias"])] for gait in gaits if gait["name"] != "straight"],
    )

    add_heading(doc, "5. Rectangle Course Controller", 1)
    doc.add_paragraph(
        "The rectangle-course viewer does not directly switch between the saved fixed gaits. "
        "Instead, it computes a steering command from the path-following controller, smooths it, "
        "then converts it into joint_bias and a tail-amplitude modulation."
    )
    doc.add_paragraph("steer_state <- steer_state + beta (target_steer - steer_state)")
    doc.add_paragraph("joint_bias_i = w_i steer_state")
    doc.add_paragraph("w = [0.45, 0.55, 0.68, 0.80, 0.92, 1.00]")
    doc.add_paragraph(
        "The default smoothing beta is 0.08. Smaller beta gives smoother but slower turning response."
    )

    base = tuple(float(v) for v in straight["amp_scales"])
    example_steer = 0.18
    turn_gain = 0.6
    turned = turning_amp_scales(base, example_steer, turn_gain)
    mu_scales = amp_scales_to_mu_scales(turned)
    add_table(
        doc,
        ["Case", "amp_scales", "mu_scales"],
        [
            ["Straight/base", fmt(base), fmt(amp_scales_to_mu_scales(base))],
            [f"Turning example steer={example_steer:.2f}, gain={turn_gain:.1f}", fmt(turned), fmt(mu_scales)],
        ],
    )
    doc.add_paragraph(
        "The turning amplitude modulation is applied through Hopf mu_i targets. Because r_i is a "
        "Hopf state, the amplitude moves toward the new value smoothly instead of jumping instantly."
    )

    add_heading(doc, "6. Practical Notes", 1)
    notes = [
        "Direct sine control is simple for a fixed gait, but changing wavelength or amplitude during swimming can create abrupt target-angle changes.",
        "Hopf CPG keeps internal r_i and theta_i states, so online gait changes are continuous and smoother.",
        "For the current model, phase lag is fixed during rectangle turning; turning mainly uses joint_bias plus tail amplitude modulation.",
        "The current slower Hopf gains make gait transitions smoother, but the gait takes longer to converge after startup or after a turn command changes.",
    ]
    for note in notes:
        doc.add_paragraph(note, style="List Bullet")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
